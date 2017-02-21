from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

import os
import sys

from Sifter import getEventNames

file_path = os.path.dirname(os.path.realpath(__file__)) + "\\" # Obtains the scripts file path
storage_dir = file_path + "EVENT_LISTS\\"

# A simple print with a system flush afterwards
def printFlush(string = ''):
	print(string)
	sys.stdout.flush()

def createDocx(gEvent, sEvent):
	import errno
	try:
		os.makedirs(storage_dir)
		printFlush("**Event list directory created!")
	except OSError as exception:
		if exception.errno != errno.EEXIST:
			raise
		else:
			printFlush("**Event list directory exists!")

	createGlobalDocx(gEvent)
	createSpecificDocx(sEvent)


def createGlobalDocx(gEvent):
	globalDoc = Document()
	globalDoc.add_heading(gEvent, 0)

	globalTxt = open(file_path + gEvent + ".txt", 'r')

	for lineNum, line in enumerate(globalTxt.readlines()) :
		if lineNum % 2 == 0: 
			paragraph = globalDoc.add_paragraph(line)
			paragraph.style = 'ListNumber'

	globalTxt.close()
	globalDoc.save(storage_dir + gEvent + ".docx")

def createSpecificDocx(sEvent):
	specificDoc = Document()
	specificDoc.add_heading(sEvent, 0)

	specificTxt = open(file_path + sEvent + ".txt", 'r')

	obj_styles = specificDoc.styles
	obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
	obj_font = obj_charstyle.font
	obj_font.size = Pt(18)

	for line in specificTxt.readlines() :
		if line.strip() != "":
			if line[:2] != "--":
				paragraph = specificDoc.add_paragraph(line)
				paragraph.style = 'ListNumber'
			else:
				paragraph = specificDoc.add_paragraph("")
				run = paragraph.add_run(line, style = 'CommentsStyle').bold = True
				paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

	specificTxt.close()
	
	specificDoc.save(storage_dir + sEvent + ".docx")

def main():
	nameset = getEventNames()
	createDocx(nameset[0], nameset[1])


if __name__ == "__main__":
	main()